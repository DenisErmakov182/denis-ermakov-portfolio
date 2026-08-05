from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

OUT = "tmp/resume-revolut/Denis Ermakov.docx"

# Design preset: compact_reference_guide.
# Named resume override: 0.65in margins and 10.5pt body text keep this one-page
# application resume concise while preserving readable spacing.
doc = Document()
section = doc.sections[0]
section.top_margin = Inches(0.65)
section.bottom_margin = Inches(0.65)
section.left_margin = Inches(0.70)
section.right_margin = Inches(0.70)
section.header_distance = Inches(0.3)
section.footer_distance = Inches(0.3)

styles = doc.styles
normal = styles["Normal"]
normal.font.name = "Arial"
normal._element.rPr.rFonts.set(qn("w:ascii"), "Arial")
normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Arial")
normal.font.size = Pt(10.5)
normal.font.color.rgb = RGBColor(31, 31, 31)
normal.paragraph_format.space_after = Pt(3)
normal.paragraph_format.line_spacing = 1.08

for name, size, before, after in [
    ("Resume Heading", 10.5, 9, 4),
    ("Resume Body", 10.5, 0, 3),
    ("Resume Meta", 9.5, 0, 2),
]:
    style = styles.add_style(name, WD_STYLE_TYPE.PARAGRAPH)
    style.base_style = normal
    style.font.name = "Arial"
    style._element.rPr.rFonts.set(qn("w:ascii"), "Arial")
    style._element.rPr.rFonts.set(qn("w:hAnsi"), "Arial")
    style.font.size = Pt(size)
    style.paragraph_format.space_before = Pt(before)
    style.paragraph_format.space_after = Pt(after)
    style.paragraph_format.line_spacing = 1.08

heading = styles["Resume Heading"]
heading.font.bold = True
heading.font.color.rgb = RGBColor(31, 77, 120)
heading.font.all_caps = True

list_bullet = styles["List Bullet"]
list_bullet.base_style = normal
list_bullet.font.name = "Arial"
list_bullet._element.rPr.rFonts.set(qn("w:ascii"), "Arial")
list_bullet._element.rPr.rFonts.set(qn("w:hAnsi"), "Arial")
list_bullet.font.size = Pt(10.5)
list_bullet.paragraph_format.left_indent = Inches(0.38)
list_bullet.paragraph_format.first_line_indent = Inches(-0.19)
list_bullet.paragraph_format.space_after = Pt(2)
list_bullet.paragraph_format.line_spacing = 1.08

def set_font(run, size=None, bold=None, color=None):
    run.font.name = "Arial"
    run._element.rPr.rFonts.set(qn("w:ascii"), "Arial")
    run._element.rPr.rFonts.set(qn("w:hAnsi"), "Arial")
    if size:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if color:
        run.font.color.rgb = RGBColor(*color)

def line_bottom(paragraph, color="9CB8D4", sz="8"):
    p_pr = paragraph._p.get_or_add_pPr()
    borders = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), sz)
    bottom.set(qn("w:space"), "4")
    bottom.set(qn("w:color"), color)
    borders.append(bottom)
    p_pr.append(borders)

def bullet(text):
    p = doc.add_paragraph(style="List Bullet")
    p.add_run(" " + text)
    return p

def role(company, title, dates, location_or_type):
    p = doc.add_paragraph(style="Resume Body")
    p.paragraph_format.space_before = Pt(1)
    p.paragraph_format.space_after = Pt(1)
    r = p.add_run(company)
    set_font(r, bold=True)
    r = p.add_run(" | " + title)
    set_font(r, bold=True)
    p.add_run("\n")
    r = p.add_run(f"{dates} | {location_or_type}")
    set_font(r, size=9.5, color=(88, 88, 88))

# Header: the filename is intentionally neutral, without ATS or a role title.
p = doc.add_paragraph()
p.paragraph_format.space_after = Pt(1)
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("DENIS ERMAKOV")
set_font(r, size=20, bold=True, color=(11, 37, 69))

p = doc.add_paragraph(style="Resume Meta")
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_after = Pt(5)
r = p.add_run("Moscow, Russia | +7 999 030 8563 | @mrakov182 | denis.ermakov@me.com")
set_font(r, size=9.5, color=(68, 68, 68))
p.add_run("\n")
r = p.add_run("denisermakov182.github.io/denis-ermakov-portfolio | linkedin.com/in/denis-ermakov-a548a0388")
set_font(r, size=9.5, color=(68, 68, 68))
line_bottom(p)

p = doc.add_paragraph(style="Resume Body")
p.paragraph_format.space_before = Pt(7)
p.paragraph_format.space_after = Pt(2)
r = p.add_run("Product Designer with 3+ years of full-time experience in EdTech and early-stage products. ")
set_font(r, bold=True)
p.add_run("I turn research findings into user flows, prototypes, and shipped features, working closely with developers and stakeholders. My experience includes qualitative and quantitative research, usability testing, AI-enabled product features, and design systems.")

doc.add_paragraph("Experience", style="Resume Heading")

role("CO Koalitsiya", "Product Designer", "Mar 2023 – Present", "Moscow | Full-time")
bullet("Conducted teacher interviews and identified key friction in publishing video lessons: reliance on third-party services, manual timecodes, and missing lesson summaries.")
bullet("Designed an end-to-end MP4 upload and AI processing flow in the lesson editor, including transcription, timecodes, and summaries; reduced time to publish by 40%, increased teacher CSAT by 18%, and reached 64% adoption of the AI feature.")
bullet("Redesigned teacher and student dashboards, designed an AI tutor for test generation, and maintained Figma components, variants, and variables with developers through delivery.")

role("GPtest", "Independent Product Designer", "Mar 2026 – Present", "Product in production")
bullet("Ran a survey in student communities with 127 responses in two weeks, validating the problem: 59% used ChatGPT and 40% lacked explanations for errors.")
bullet("Launched an MVP to 20 users, collected daily feedback for two weeks, and iterated through usability testing with four users.")
bullet("Improved the testing flow with progress analytics, accidental-close protection, clearer labels, file names, and a generation loading state.")

doc.add_paragraph("Selected projects", style="Resume Heading")
role("Mappy", "Independent Product Designer", "Jul 2025 – Sep 2025", "Product launched to production")
bullet("Built a geo-service for saving places: surveyed 57 respondents, formulated four core scenario hypotheses, validated them in two rounds of hallway testing, and created the UI kit and iterative product layouts.")

doc.add_paragraph("Education", style="Resume Heading")
p = doc.add_paragraph(style="Resume Body")
p.paragraph_format.space_after = Pt(2)
r = p.add_run("International Banking Institute named after Anatoliy Sobchak")
set_font(r, bold=True)
p.add_run("\nBachelor's degree in Business Informatics, expected 2029 | Online programme, part-time")

doc.add_paragraph("Languages", style="Resume Heading")
p = doc.add_paragraph(style="Resume Body")
p.paragraph_format.space_after = Pt(0)
p.add_run("Russian — Native | English — B2 (Upper-Intermediate)")

doc.core_properties.title = "Denis Ermakov"
doc.core_properties.subject = "Resume"
doc.core_properties.author = "Denis Ermakov"
doc.save(OUT)
